import shutil
import sys
import types
import unittest
from pathlib import Path

from docx import Document

# Make DOCX generator import independent from runtime-only settings dependency.
if "app.config" not in sys.modules:
    fake_config = types.ModuleType("app.config")
    fake_config.settings = types.SimpleNamespace(REPORTS_DIR=".")
    sys.modules["app.config"] = fake_config

from app.reports.docx_generator import DOCXGenerator


def _doc_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text:
                        parts.append(paragraph.text)
    return "\n".join(parts)


class RemainingDocxToolsTests(unittest.TestCase):
    def setUp(self):
        self.temp_dir = Path("tests") / ".tmp_docx_remaining"
        if self.temp_dir.exists():
            shutil.rmtree(self.temp_dir)
        self.temp_dir.mkdir(parents=True, exist_ok=True)
        self.generator = DOCXGenerator()
        self.generator.reports_dir = str(self.temp_dir)

    def tearDown(self):
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def _build_and_text(self, method_name: str, task_id: str, data: dict) -> str:
        report_path = getattr(self.generator, method_name)(task_id, data)
        return _doc_text(Document(report_path))

    def test_robots_docx_contains_appendix(self):
        text = self._build_and_text(
            "generate_robots_report",
            "robots-docx-smoke",
            {
                "url": "https://example.com",
                "results": {
                    "robots_txt_found": True,
                    "status_code": 200,
                    "quality_score": 91,
                    "quality_grade": "A",
                    "production_ready": True,
                    "quick_status": "ok",
                    "content_length": 120,
                    "lines_count": 8,
                    "user_agents": 1,
                    "disallow_rules": 1,
                    "allow_rules": 1,
                    "sitemaps": ["https://example.com/sitemap.xml"],
                    "hosts": ["example.com"],
                    "crawl_delays": {},
                    "clean_params": [],
                    "severity_counts": {"critical": 0, "warning": 1, "info": 1},
                    "warning_issues": ["Missing explicit GPTBot rule"],
                    "info_issues": ["Robots file is reachable"],
                    "recommendations": ["Add explicit AI bot policy."],
                    "ui_only_block": {"widget_title": "Robots UI card"},
                },
            },
        )
        self.assertIn("Robots.txt Report", text)
        self.assertIn("Приложение. Полная структура результата", text)
        self.assertIn("Robots UI card", text)

    def test_sitemap_docx_contains_appendix(self):
        text = self._build_and_text(
            "generate_sitemap_report",
            "sitemap-docx-smoke",
            {
                "url": "https://example.com",
                "results": {
                    "input_url": "https://example.com",
                    "resolved_sitemap_url": "https://example.com/sitemap.xml",
                    "sitemap_discovery_source": "robots.txt",
                    "valid": True,
                    "status_code": 200,
                    "sitemaps_scanned": 1,
                    "sitemaps_valid": 1,
                    "urls_count": 12,
                    "unique_urls_count": 12,
                    "duplicate_urls_count": 0,
                    "quality_score": 95,
                    "quality_grade": "A",
                    "severity_counts": {"critical": 0, "warning": 0, "info": 1},
                    "issues": [],
                    "recommendations": ["Keep sitemap fresh."],
                    "ui_only_snapshot": {"coverage_badge": "100% indexed sample"},
                },
            },
        )
        self.assertIn("Sitemap Report", text)
        self.assertIn("Приложение. Полная структура результата", text)
        self.assertIn("100% indexed sample", text)

    def test_mobile_docx_contains_appendix(self):
        text = self._build_and_text(
            "generate_mobile_report",
            "mobile-docx-smoke",
            {
                "url": "https://example.com",
                "results": {
                    "engine": "legacy",
                    "mode": "full",
                    "score": 82,
                    "mobile_friendly": False,
                    "status_code": 200,
                    "final_url": "https://example.com",
                    "viewport_found": True,
                    "viewport_content": "width=device-width, initial-scale=1",
                    "summary": {
                        "total_devices": 1,
                        "mobile_friendly_devices": 0,
                        "non_friendly_devices": 1,
                        "avg_load_time_ms": 1450,
                    },
                    "device_results": [
                        {
                            "device_name": "iPhone 14",
                            "category": "phone",
                            "viewport": {"width": 390, "height": 844},
                            "status_code": 200,
                            "load_time_ms": 1450,
                            "issues_count": 1,
                            "mobile_friendly": False,
                            "issues": [
                                {
                                    "code": "small_touch_targets",
                                    "severity": "warning",
                                    "title": "Small touch targets",
                                    "details": "Buttons are too close.",
                                    "device": "iPhone 14",
                                }
                            ],
                        }
                    ],
                    "issues": [
                        {
                            "code": "small_touch_targets",
                            "severity": "warning",
                            "title": "Small touch targets",
                            "details": "Buttons are too close.",
                            "device": "iPhone 14",
                        }
                    ],
                    "ui_only_panel": {"summary": "Mobile dashboard card"},
                },
            },
        )
        self.assertIn("Mobile Audit Report", text)
        self.assertIn("Приложение. Полная структура результата", text)
        self.assertIn("Mobile dashboard card", text)

    def test_render_docx_contains_appendix(self):
        text = self._build_and_text(
            "generate_render_report",
            "render-docx-smoke",
            {
                "url": "https://example.com/page",
                "results": {
                    "engine": "legacy",
                    "summary": {
                        "variants_total": 1,
                        "score": 78,
                        "critical_issues": 1,
                        "warning_issues": 1,
                        "missing_total": 3,
                        "avg_missing_pct": 12.5,
                    },
                    "variants": [
                        {
                            "variant_id": "desktop",
                            "variant_label": "Desktop",
                            "profile_type": "desktop",
                            "metrics": {"score": 78, "total_missing": 3, "missing_pct": 12.5},
                            "raw": {
                                "title": "No JS title",
                                "meta_description": "No JS description",
                                "h1_count": 1,
                                "h2_count": 2,
                                "images_count": 4,
                                "links_count": 6,
                                "canonical": "https://example.com/page",
                                "structured_data_count": 0,
                            },
                            "rendered": {
                                "title": "JS title",
                                "meta_description": "JS description",
                                "h1_count": 1,
                                "h2_count": 3,
                                "images_count": 4,
                                "links_count": 8,
                                "canonical": "https://example.com/page",
                                "structured_data_count": 1,
                                "schema_types": ["Article"],
                            },
                            "meta_non_seo": {
                                "raw": {"meta:viewport": "width=device-width"},
                                "rendered": {"meta:viewport": "width=device-width"},
                                "comparison": {"items": []},
                            },
                            "issues": [{"severity": "critical", "title": "JS-only text", "details": "Content appears only after JS."}],
                            "recommendations": ["SSR important text."],
                        }
                    ],
                    "issues": [{"severity": "critical", "variant": "Desktop", "title": "JS-only text", "details": "Content appears only after JS."}],
                    "recommendations": ["SSR important text."],
                    "ui_only_render_card": {"delta": "JS gap visible"},
                },
            },
        )
        self.assertIn("Render Audit Report", text)
        self.assertIn("Приложение. Полная структура результата", text)
        self.assertIn("JS gap visible", text)

    def test_onpage_docx_contains_appendix(self):
        text = self._build_and_text(
            "generate_onpage_report",
            "onpage-docx-smoke",
            {
                "url": "https://example.com/page",
                "results": {
                    "engine": "onpage-v2",
                    "score": 74,
                    "status_code": 200,
                    "final_url": "https://example.com/page",
                    "language": "ru",
                    "summary": {
                        "score": 74,
                        "spam_score": 12,
                        "keyword_coverage_score": 68,
                        "keyword_coverage_pct": 44,
                        "ai_risk_composite": 21,
                        "critical_issues": 1,
                        "warning_issues": 2,
                        "info_issues": 1,
                    },
                    "content": {"word_count": 850, "unique_word_count": 420, "char_count": 5400},
                    "content_profile": {"clean_text_length": 5000, "core_vocabulary": 230, "wateriness_pct": 11, "nausea_index": 3.2, "text_html_pct": 32},
                    "title": {"text": "Example Title", "length": 13},
                    "description": {"text": "Example Description", "length": 19},
                    "h1": {"count": 1, "values": ["Example H1"]},
                    "keywords": [{"keyword": "example", "occurrences": 7, "density_pct": 1.2, "in_title": True, "in_description": True, "in_h1": True, "status": "ok"}],
                    "top_terms": [{"term": "example", "count": 7, "pct": 1.2}],
                    "technical": {"canonical_href": "https://example.com/page", "canonical_is_self": True, "robots": "index,follow", "viewport": "width=device-width", "lang": "ru", "hreflang_count": 0, "schema_count": 1},
                    "links": {"links_total": 12, "internal_links": 10, "external_links": 2, "nofollow_links": 1, "empty_anchor_links": 0},
                    "media": {"images_total": 4, "images_missing_alt": 1},
                    "readability": {"sentences_count": 42, "avg_sentence_len": 13, "long_sentence_ratio": 0.1, "lexical_diversity": 0.62},
                    "spam_metrics": {"stopword_ratio": 0.12, "content_html_ratio": 0.32, "uppercase_ratio": 0.01, "punctuation_ratio": 0.08, "duplicate_sentences": 1, "duplicate_sentence_ratio": 0.02},
                    "ngrams": {"bigrams": [{"term": "example page", "count": 4, "pct": 0.5}], "trigrams": [{"term": "example page title", "count": 2, "pct": 0.2}]},
                    "schema": {"json_ld_blocks": 1, "json_ld_valid_blocks": 1, "microdata_items": 0, "rdfa_items": 0, "types": [{"type": "Article"}]},
                    "opengraph": {"tags_count": 4, "required_present_count": 4, "required_missing": []},
                    "ai_insights": {"ai_marker_density_1k": 0.2, "hedging_ratio": 0.1, "template_repetition": 0.05, "burstiness_cv": 0.4, "perplexity_proxy": 56, "entity_depth_1k": 2, "claim_specificity_score": 71, "author_signal_score": 66, "source_attribution_score": 59, "ai_risk_composite": 21},
                    "issues": [{"severity": "critical", "title": "Missing FAQ block", "details": "No FAQ section on commercial page."}],
                    "recommendations": ["Add FAQ section."],
                    "ui_only_heatmap": {"content": "hot"},
                },
            },
        )
        self.assertIn("OnPage Audit Report", text)
        self.assertIn("Приложение. Полная структура результата", text)
        self.assertIn("hot", text)

    def test_site_audit_pro_docx_contains_cover_and_appendix(self):
        text = self._build_and_text(
            "generate_site_audit_pro_report",
            "site-audit-pro-docx-smoke",
            {
                "url": "https://example.com",
                "results": {
                    "mode": "full",
                    "summary": {
                        "total_pages": 15,
                        "issues_total": 4,
                        "critical_issues": 1,
                        "warning_issues": 2,
                        "info_issues": 1,
                        "score": 81,
                    },
                    "pipeline": {
                        "metrics": {
                            "avg_response_time_ms": 820,
                            "avg_readability_score": 58,
                            "avg_link_quality_score": 74,
                            "orphan_pages": 1,
                        },
                        "pagerank": [{"url": "https://example.com/", "score": 0.81}],
                        "topic_clusters": [{"topic": "products", "count": 7, "urls": ["https://example.com/products"]}],
                    },
                    "pages": [{"url": "https://example.com/products", "recommendation": "Improve internal linking."}],
                    "issues": [{"severity": "critical", "code": "orphan_page", "url": "https://example.com/orphan", "title": "Orphan page"}],
                    "ui_only_story": {"badge": "crawl budget risk"},
                },
            },
        )
        self.assertIn("Site Audit Pro Report", text)
        self.assertIn("Приложение. Полная структура результата", text)
        self.assertIn("crawl budget risk", text)

    def test_link_profile_docx_appendix_preserves_extra_fields(self):
        text = self._build_and_text(
            "generate_link_profile_report",
            "link-profile-docx-smoke",
            {
                "url": "example.com",
                "results": {
                    "summary": {
                        "our_domain": "example.com",
                        "rows_total": 2,
                        "unique_ref_domains": 2,
                        "unique_competitors": 1,
                        "our_links": 2,
                        "dofollow": 1,
                        "nofollow": 1,
                        "unknown_follow": 0,
                        "dofollow_pct": 50,
                        "nofollow_pct": 50,
                        "lost_links_pct": 10,
                        "http_2xx_pct": 100,
                        "avg_dr": 47,
                    },
                    "tables": {
                        "competitor_benchmark": [
                            {"domain": "alpha.com", "score": 80},
                            {"domain": "beta.com", "score": 77, "extra_note": "Shown only in second row"},
                        ]
                    },
                    "warnings": ["Low domain diversity."],
                    "keywords": {"derivedBrandKeywords": ["example brand"]},
                },
            },
        )
        self.assertIn("Link Profile Report", text)
        self.assertIn("Приложение. Полная структура результата", text)
        self.assertIn("Shown only in second row", text)


if __name__ == "__main__":
    unittest.main()
